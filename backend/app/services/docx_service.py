from typing import List, Dict, Any
import os
import docx  # from python-docx

class DocxService:
    @staticmethod
    def extract_text(file_path: str) -> List[Dict[str, Any]]:
        """
        Extract text from a DOCX file with 'chunk' metadata.
        Since DOCX doesn't have strict pages, we group by paragraphs.
        
        Returns:
            List[Dict]: [{'page_content': '...', 'metadata': {'chunk_index': i}}]
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX file not found at: {file_path}")
            
        docs = []
        try:
            doc = docx.Document(file_path)
            
            current_chunk = []
            chunk_char_count = 0
            chunk_idx = 1
            TARGET_CHUNK_SIZE = 800  # Approx characters per "page" equivalent

            # Helper to flush
            def flush_chunk():
                nonlocal current_chunk, chunk_char_count, chunk_idx
                if current_chunk:
                    text = "\n".join(current_chunk).strip()
                    if text:
                        docs.append({
                            "page_content": text,
                            "metadata": {"page": chunk_idx} # Use 'page' key for consistency with PDF
                        })
                        chunk_idx += 1
                    current_chunk = []
                    chunk_char_count = 0

            # 1. Process Paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                current_chunk.append(text)
                chunk_char_count += len(text)
                
                if chunk_char_count >= TARGET_CHUNK_SIZE:
                    flush_chunk()
            
            # 2. Process Tables (treat each table as a separate chunk/page usually)
            for table in doc.tables:
                table_texts = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        table_texts.append(" | ".join(row_data))
                
                if table_texts:
                    # Flush previous content first
                    flush_chunk()
                    # Add table
                    docs.append({
                        "page_content": "\n".join(table_texts),
                        "metadata": {"page": chunk_idx}
                    })
                    chunk_idx += 1

            # Final flush
            flush_chunk()

            return docs
            
        except Exception as e:
            raise ValueError(f"Failed to read DOCX: {str(e)}")